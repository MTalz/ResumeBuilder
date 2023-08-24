# Databricks notebook source
# DBTITLE 1,Import relevant libraries
pip install python-docx -q


# COMMAND ----------

dbutils.library.restartPython()

# COMMAND ----------

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import pandas as pd

# COMMAND ----------

# MAGIC %run "/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Utility/Utility Notebook"

# COMMAND ----------

# DBTITLE 1,Import resume components
# Name, Email, Phone Number, LinkedIn
UserInfoPath = "/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Experience History/UserInfo.csv"
with open(UserInfoPath) as f:
    lines = f.read().splitlines()

## Set these values to variables
name = lines[0].split(',')[1].strip()
email = lines[1].split(',')[1].strip()
mobile = lines[2].split(',')[1].strip()
linkedin = lines[3].split(',')[1].strip()

# COMMAND ----------

# Import work history roles as df
wh_roles = pd.read_csv("/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Experience History/WorkHistory_Roles.csv")


# COMMAND ----------

# Obtain job dictionaries
job_ads = obtain_ad_folder_dict()

# COMMAND ----------

# DBTITLE 1,Generate Resume
for job, folder in job_ads.items():

    # Higher Education
    higher_ed_path = f"/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Resume Components/{folder}/HigherEd.txt"
    with open(higher_ed_path) as f:
        higher_ed = f.read().splitlines()
    f.close()

    # Other Education
    other_ed_path = f"/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Resume Components/{folder}/Rel_OtherEd.txt"
    with open(other_ed_path) as f:
        other_ed = f.read().splitlines()
    f.close()

    # Hard Skills
    hard_skills_path = f"/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Resume Components/{folder}/Rel_Skills.txt"
    with open(hard_skills_path) as f:
        hard_skills = f.read().splitlines()
    f.close()

    # Relevant achievements
    # Read in achievements
    wh_path =   f'/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Resume Components/{folder}/sorted_achievements.csv'
    wh_sorted_df = pd.read_csv(wh_path)

    # Create an list version of achievments
    wh_sorted_df['achievement_list'] = wh_sorted_df['achievement'].apply( lambda x: [x])

    # Create grouped list of achievements
    wh_grouped = wh_sorted_df.groupby('role_id').agg({ 'achievement_list' : 'sum'})

    # Join to wh_roles
    relevant_wh_df = wh_grouped.merge(wh_roles, how  = 'left', on = 'role_id' )

    # Generate and save Resume
    document = Document()

    # Set Margins
    sections = document.sections
    section = sections[0] # Access the first section
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.2)
    section.botom_margin = Inches(0.5)

    # Centre Name at top of Resume
    heading = document.add_heading(f'{name}', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Put in contact details
    contact_table = document.add_table(rows=0, cols=3)
    contact_row = contact_table.add_row().cells
    p = contact_row[0].paragraphs[0] # Put this in to remove the first paragraph
    p.add_run('Email:\n')
    p.add_run(f'{email}').bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = contact_row[1].paragraphs[0]
    p.add_run('Mobile:\n')
    p.add_run(f"{mobile}").bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = contact_row[2].paragraphs[0]
    p.add_run('LinkedIn:\n')
    p.add_run(f"{linkedin}").bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Change font in entire table
    for row in contact_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(8)

    # Put in Work Experience heading
    heading_3 = document.add_paragraph('Work Experience', style='Intense Quote')
    heading_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ## Work Experience Table
    we_table = document.add_table(rows=0, cols=2)

    for index, row in relevant_wh_df.iterrows():
        we_row = we_table.add_row().cells
        # #### Role Dates
        we_row[0].width = Inches(1) # Set the cell width
        p = we_row[0].paragraphs[0] # Put this in to remove the first paragraph
        p.add_run(row['year'])

        we_row[1].width = Inches(7.5) # Set the cell width
        p = we_row[1].paragraphs[0]
        p.add_run(f"{row['role']}\n").bold =True
        p.add_run(f"{row['company']}\n").italic = True

        for achievement in row['achievement_list']:
            p.add_run(f"- {achievement}\n")

    # Put in Higher Education
    heading_2 = document.add_paragraph('Education', style='Intense Quote')
    for item in higher_ed:
        document.add_paragraph(item, style='List Bullet') #Cycle through higher education and list as bullet points

    # Put in Other Education
    heading_3 = document.add_paragraph('Other Certifications', style='Intense Quote')
    for item in other_ed:
        document.add_paragraph(item, style='List Bullet') #Cycle through higher education and list as bullet points

    # Put in Other Education
    heading_4 = document.add_paragraph('Relevant Skills', style='Intense Quote')
    for item in hard_skills:
        document.add_paragraph(item, style='List Bullet') #Cycle through higher education and list as bullet points


    # References provided upon request
    p = document.add_paragraph()
    p.add_run('References provided upon request').bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # document.save(f'/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Output/{folder}/Resume.docx')
    document.save(f'/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Output/{folder}/Resume.docx')


    # Generate and save cover letter
    # Cover Letter Text
    cover_letter_text_path = f"/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Resume Components/{folder}/cover_letter_text.txt"
    with open(cover_letter_text_path) as f:
        cover_letter_text = f.read()
    f.close()


    cover_letter = Document()

    # Centre Name at top of Resume
    heading = cover_letter.add_heading(f'{name}', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Put in contact details
    contact_table = cover_letter.add_table(rows=0, cols=3)
    contact_row = contact_table.add_row().cells
    p = contact_row[0].paragraphs[0] # Put this in to remove the first paragraph
    p.add_run('Email:\n')
    p.add_run(f'{email}').bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = contact_row[1].paragraphs[0]
    p.add_run('Mobile:\n')
    p.add_run(f"{mobile}").bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = contact_row[2].paragraphs[0]
    p.add_run('LinkedIn:\n')
    p.add_run(f"{linkedin}").bold =True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Put in cover letter text
    p = cover_letter.add_paragraph(cover_letter_text)

    
    cover_letter.save(f'/Workspace/Users/joe.g.gulay@au.ey.com/Stonecutters/Resume Builder/Output/{folder}/Cover_Letter.docx')


# COMMAND ----------

